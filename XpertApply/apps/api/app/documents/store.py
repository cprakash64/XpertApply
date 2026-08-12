"""Persistence, serialization, and export for generated documents."""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from app.core.audit import record_audit
from app.models.entities import DocumentFormat, DocumentType, GeneratedDocument, JobPosting
from app.services.documents import profile_payload, public_dict


def persist_document(
    db: Session,
    user_id: int,
    job: JobPosting,
    document_type: DocumentType,
    *,
    title: str,
    content: dict[str, Any],
    markdown: str,
    plain_text: str = "",
    quality: dict[str, Any] | None = None,
    model_used: str,
) -> GeneratedDocument:
    record = GeneratedDocument(
        user_id=user_id,
        job_id=job.id,
        type=document_type,
        format=DocumentFormat.json,
        title=title,
        content=_json_safe(content),
        content_markdown=markdown,
        plain_text=plain_text,
        quality=_json_safe(quality or {}),
        source_profile_snapshot=_json_safe(profile_payload(db, user_id)),
        job_snapshot=_json_safe(public_dict(job)),
        model_used=model_used,
        format_version="v2",
    )
    db.add(record)
    record_audit(db, user_id, "document_generated", {"job_id": job.id, "type": document_type.value})
    db.commit()
    db.refresh(record)
    return record


def serialize_document(
    record: GeneratedDocument,
    *,
    warnings: list[str] | None = None,
    unsupported_claims_removed: list[str] | None = None,
) -> dict[str, Any]:
    quality = record.quality or {}
    return {
        "document_id": record.id,
        "document_type": record.type.value,
        "title": record.title,
        "content": record.content,
        "markdown": record.content_markdown or "",
        "plain_text": record.plain_text or "",
        "quality": quality,
        "model_used": record.model_used,
        # Kept for backward compatibility; the frontend now reads quality.warnings.
        "warnings": warnings if warnings is not None else quality.get("warnings", []),
        "unsupported_claims_removed": (
            unsupported_claims_removed
            if unsupported_claims_removed is not None
            else quality.get("unsupported_claims_removed", [])
        ),
        "download_urls": {
            "docx": f"/jobs/documents/{record.id}/download/docx",
            "pdf": f"/jobs/documents/{record.id}/download/pdf",
        },
    }


def export_document(record: GeneratedDocument, fmt: DocumentFormat) -> str:
    out_dir = Path(os.getenv("UPLOAD_DIR", "uploads")).parent / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"document-{record.id}.{fmt.value}"
    is_resume = record.type == DocumentType.resume
    content = (
        _normalize_resume_content(record.content or {})
        if is_resume
        else (record.content or {})
    )
    if fmt == DocumentFormat.docx:
        _render_docx(content, is_resume, path)
        record.docx_file_path = str(path)
    elif fmt == DocumentFormat.pdf:
        _render_pdf(content, is_resume, path)
        record.pdf_file_path = str(path)
    else:
        path.write_text(record.plain_text or "", encoding="utf-8")
    return str(path)


# --------------------------------------------------------------------------- #
# DOCX rendering from structured content (ATS-friendly, matches the preview)
# --------------------------------------------------------------------------- #
def _render_docx(content: dict[str, Any], is_resume: bool, path: Path) -> None:
    doc = Document()
    if is_resume:
        _render_resume_docx(doc, content)
    else:
        _render_cover_docx(doc, content)
    doc.save(path)


def _render_resume_docx(doc, content: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    navy = RGBColor(25, 45, 74)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    # Word otherwise substitutes a theme font in some viewers even though the
    # visible Latin font says Arial. Set every script family so DOCX and PDF
    # keep the same ATS-safe typeface.
    normal.element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    def add_section(title: str) -> None:
        # Put one clean text-line of breathing room between sections. Applying
        # it to the final content paragraph avoids empty paragraphs (which some
        # ATS parsers expose as blank records) and still lets Word paginate
        # naturally to a second page only when the content genuinely needs it.
        if doc.paragraphs:
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.widow_control = True
        run = paragraph.add_run(title.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = navy
        # A paragraph border is real Word formatting (not an image/table), so
        # it matches the PDF/preview while remaining simple for ATS extraction.
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "5")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "9DB6D3")
        borders.append(bottom)
        properties.append(borders)

    def add_bullet(text: str) -> None:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(0.8)
        paragraph.paragraph_format.widow_control = True
        paragraph.add_run(str(text))

    header = content.get("header") or {}
    if header.get("full_name"):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(str(header["full_name"]).upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(17)
        run.font.color.rgb = navy
    contact = " | ".join(
        str(v)
        for v in [
            header.get("email"),
            header.get("phone"),
            header.get("location"),
            *(header.get("links") or []),
        ]
        if v
    )
    if contact:
        paragraph = doc.add_paragraph(contact)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            run.font.size = Pt(8.5)

    if content.get("summary"):
        add_section("Professional Summary")
        doc.add_paragraph(content["summary"])

    skills = [g for g in content.get("skills") or [] if g.get("items")]
    if skills:
        add_section("Core Skills")
        for group in skills:
            paragraph = doc.add_paragraph()
            if group.get("category"):
                paragraph.add_run(f"{group['category']}: ").bold = True
            paragraph.add_run(", ".join(str(item) for item in group["items"]))

    if content.get("experience"):
        add_section("Professional Experience")
        for exp in content["experience"]:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(2)
            head.paragraph_format.space_after = Pt(0.5)
            head.paragraph_format.keep_with_next = bool(exp.get("bullets"))
            head.paragraph_format.widow_control = True
            heading_text = " | ".join(filter(None, [exp.get("title"), exp.get("company")]))
            head.add_run(heading_text).bold = True
            meta = " | ".join(str(v) for v in [exp.get("location"), exp.get("dates")] if v)
            if meta:
                head.add_run(f" - {meta}")
            for bullet in exp.get("bullets") or []:
                add_bullet(bullet)

    if content.get("projects"):
        add_section("Selected Projects")
        for proj in content["projects"]:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(2)
            head.paragraph_format.space_after = Pt(0.5)
            head.paragraph_format.keep_with_next = bool(proj.get("bullets"))
            head.paragraph_format.widow_control = True
            head.add_run(proj.get("name", "")).bold = True
            tech = ", ".join(proj.get("technologies") or [])
            if tech:
                head.add_run(f" | {tech}")
            for bullet in proj.get("bullets") or []:
                add_bullet(bullet)

    if content.get("education"):
        add_section("Education")
        for edu in content["education"]:
            paragraph = doc.add_paragraph()
            if edu.get("school"):
                paragraph.add_run(str(edu["school"])).bold = True
            remainder = " | ".join(
                str(v)
                for v in [edu.get("degree"), edu.get("dates"), edu.get("details")]
                if v
            )
            if remainder:
                paragraph.add_run(f" | {remainder}" if edu.get("school") else remainder)

    awards = (content.get("awards") or []) + (content.get("certifications") or [])
    awards = [a for a in awards if (a.get("name") if isinstance(a, dict) else a)]
    if awards:
        add_section("Awards & Certifications")
        for award in awards:
            name = award.get("name") if isinstance(award, dict) else str(award)
            add_bullet(name)


def _render_cover_docx(doc, content: dict[str, Any]) -> None:
    doc.add_paragraph(content.get("date", ""))
    doc.add_paragraph(content.get("recipient", "Hiring Team"))
    if content.get("company"):
        doc.add_paragraph(content["company"])
    doc.add_paragraph("")
    doc.add_paragraph(content.get("greeting", "Dear Hiring Team,"))
    for para in content.get("paragraphs") or []:
        doc.add_paragraph(para)
    doc.add_paragraph("")
    doc.add_paragraph(content.get("closing", "Best regards,"))
    doc.add_paragraph(content.get("signature", ""))


def _render_pdf(content: dict[str, Any], is_resume: bool, path: Path) -> None:
    """Render structured, selectable, single-column ATS-safe PDF content."""
    from xml.sax.saxutils import escape

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    navy = colors.HexColor("#192D4A")
    muted = colors.HexColor("#334155")
    body = ParagraphStyle(
        "ResumeBody", fontName="Helvetica", fontSize=8.9, leading=10.5,
        textColor=colors.black, spaceAfter=1.5,
    )
    centered = ParagraphStyle(
        "ResumeContact", parent=body, alignment=TA_CENTER, fontSize=8.2,
        leading=9.6, textColor=muted, spaceAfter=4,
    )
    name = ParagraphStyle(
        "ResumeName", parent=body, alignment=TA_CENTER, fontName="Helvetica-Bold",
        fontSize=17, leading=19, textColor=navy, spaceAfter=1,
    )
    heading = ParagraphStyle(
        "ResumeSection", parent=body, fontName="Helvetica-Bold", fontSize=10.3,
        leading=11.5, textColor=navy, spaceBefore=1, spaceAfter=1.5,
        keepWithNext=True,
    )
    entry = ParagraphStyle(
        "ResumeEntry", parent=body, leading=10.1, spaceBefore=1.5, spaceAfter=0.5,
    )
    entry_with_detail = ParagraphStyle(
        "ResumeEntryWithDetail", parent=entry, keepWithNext=True,
    )
    bullet = ParagraphStyle(
        "ResumeBullet", parent=body, leftIndent=11, firstLineIndent=-7,
        bulletIndent=4, spaceAfter=0.8,
    )

    doc = SimpleDocTemplate(
        str(path), pagesize=letter, rightMargin=0.56 * inch, leftMargin=0.56 * inch,
        topMargin=0.48 * inch, bottomMargin=0.48 * inch,
        title=str((content.get("header") or {}).get("full_name") or "Resume"),
        author=str((content.get("header") or {}).get("full_name") or ""),
        subject="ATS-friendly professional resume" if is_resume else "Cover letter",
        pageCompression=1,
    )
    story: list[Any] = []
    section_started = False

    def text(value: Any) -> str:
        return escape(str(value or "")).replace("–", "-").replace("—", "-")

    def section_title(title: str) -> None:
        nonlocal section_started
        if section_started:
            # Roughly one body-text line between the preceding section and the
            # next heading. This deliberately uses real layout space rather
            # than blank text so extraction remains clean.
            story.append(Spacer(1, 6))
        section_started = True
        story.append(Paragraph(text(title.upper()), heading))
        rule = HRFlowable(
            width="100%",
            thickness=0.65,
            color=colors.HexColor("#9DB6D3"),
            spaceAfter=2.4,
        )
        # Keep the rule attached to the first line of section content. This
        # prevents an orphaned heading at the foot of a page.
        rule.keepWithNext = True
        story.append(rule)

    def add_bullets(values: list[Any]) -> None:
        for value in values:
            if value:
                story.append(Paragraph(text(value), bullet, bulletText="•"))

    if is_resume:
        header = content.get("header") or {}
        if header.get("full_name"):
            story.append(Paragraph(text(header["full_name"]).upper(), name))
        contact_values = [
            header.get("email"), header.get("phone"), header.get("location"),
            *(header.get("links") or []),
        ]
        contact = " | ".join(text(value) for value in contact_values if value)
        if contact:
            story.append(Paragraph(contact, centered))

        if content.get("summary"):
            section_title("Professional Summary")
            story.append(Paragraph(text(content["summary"]), body))

        skills = [group for group in content.get("skills") or [] if group.get("items")]
        if skills:
            section_title("Core Skills")
            for group in skills:
                category = (
                    f"<b>{text(group.get('category'))}:</b> "
                    if group.get("category")
                    else ""
                )
                items = ", ".join(text(item) for item in group["items"])
                story.append(Paragraph(category + items, body))

        if content.get("experience"):
            section_title("Professional Experience")
            for item in content["experience"]:
                primary = " | ".join(text(v) for v in [item.get("title"), item.get("company")] if v)
                meta = " | ".join(text(v) for v in [item.get("location"), item.get("dates")] if v)
                line = f"<b>{primary}</b>" + (f" - {meta}" if meta else "")
                bullets = item.get("bullets") or []
                story.append(Paragraph(line, entry_with_detail if bullets else entry))
                add_bullets(bullets)

        if content.get("projects"):
            section_title("Selected Projects")
            for item in content["projects"]:
                technologies = ", ".join(text(v) for v in item.get("technologies") or [])
                line = f"<b>{text(item.get('name'))}</b>"
                if technologies:
                    line += f" | {technologies}"
                bullets = item.get("bullets") or []
                story.append(Paragraph(line, entry_with_detail if bullets else entry))
                add_bullets(bullets)

        if content.get("education"):
            section_title("Education")
            for item in content["education"]:
                school = text(item.get("school"))
                remainder = " | ".join(
                    text(v)
                    for v in [item.get("degree"), item.get("dates"), item.get("details")]
                    if v
                )
                line = f"<b>{school}</b>" if school else ""
                if remainder:
                    line += f" | {remainder}" if school else remainder
                story.append(Paragraph(line, entry))

        awards = (content.get("awards") or []) + (content.get("certifications") or [])
        award_names = [item.get("name") if isinstance(item, dict) else item for item in awards]
        award_names = [value for value in award_names if value]
        if award_names:
            section_title("Awards & Certifications")
            add_bullets(award_names)
    else:
        story.append(Paragraph(text(content.get("date")), body))
        story.append(Spacer(1, 7))
        story.append(Paragraph(text(content.get("recipient") or "Hiring Team"), body))
        if content.get("company"):
            story.append(Paragraph(text(content["company"]), body))
        story.append(Spacer(1, 9))
        story.append(Paragraph(text(content.get("greeting") or "Dear Hiring Team,"), body))
        story.append(Spacer(1, 4))
        for paragraph in content.get("paragraphs") or []:
            story.append(Paragraph(text(paragraph), body))
            story.append(Spacer(1, 5))
        story.append(Spacer(1, 5))
        story.append(Paragraph(text(content.get("closing") or "Best regards,"), body))
        story.append(Paragraph(text(content.get("signature")), body))

    doc.build(story)


def _normalize_resume_content(content: dict[str, Any]) -> dict[str, Any]:
    """Normalize both current and legacy resume records for one export path.

    Older ``/jobs/{id}/documents/resume`` records used ``header.name`` and a
    flat list of skills. Without normalizing them, the legacy export endpoint
    either fell back to a raw dictionary dump or could not use the structured
    renderer at all. Returning a fresh dict also avoids mutating the audited
    document snapshot during a download.
    """
    header = content.get("header") if isinstance(content.get("header"), dict) else {}
    raw_skills = content.get("skills") if isinstance(content.get("skills"), list) else []
    if raw_skills and all(not isinstance(item, dict) for item in raw_skills):
        skills = [{"category": "Core Skills", "items": [str(item) for item in raw_skills if item]}]
    else:
        skills = [
            {
                "category": str(group.get("category") or ""),
                "items": [str(item) for item in group.get("items") or [] if item],
            }
            for group in raw_skills
            if isinstance(group, dict)
        ]

    def entries(key: str) -> list[dict[str, Any]]:
        values = content.get(key)
        if not isinstance(values, list):
            return []
        return [dict(item) for item in values if isinstance(item, dict)]

    education = []
    for item in entries("education"):
        degree = str(item.get("degree") or "")
        major = str(item.get("major") or "")
        if major and major.lower() not in degree.lower():
            degree = ", ".join(filter(None, [degree, f"in {major}"]))
        dates = str(item.get("dates") or "")
        if not dates:
            start, end = item.get("start_date"), item.get("end_date")
            dates = " - ".join(str(value) for value in (start, end) if value)
        details = str(item.get("details") or "")
        if not details and item.get("gpa"):
            details = f"GPA {item['gpa']}"
        education.append(
            {
                "school": str(item.get("school") or ""),
                "degree": degree,
                "dates": dates,
                "details": details,
            }
        )

    def named_entries(key: str) -> list[dict[str, Any]]:
        values = content.get(key)
        if not isinstance(values, list):
            return []
        return [
            {"name": str(item.get("name") or "")} if isinstance(item, dict) else {"name": str(item)}
            for item in values
            if (item.get("name") if isinstance(item, dict) else item)
        ]

    return {
        "header": {
            "full_name": str(header.get("full_name") or header.get("name") or ""),
            "email": str(header.get("email") or ""),
            "phone": str(header.get("phone") or ""),
            "location": str(header.get("location") or ""),
            "links": [str(value) for value in header.get("links") or [] if value],
        },
        "summary": str(content.get("summary") or ""),
        "skills": skills,
        "experience": entries("experience"),
        "projects": entries("projects"),
        "education": education,
        "awards": named_entries("awards"),
        "certifications": named_entries("certifications"),
    }


def _json_safe(value: Any) -> Any:
    """Recursively convert datetimes/dates to strings so snapshots persist as JSON."""
    return json.loads(json.dumps(value, default=str))
