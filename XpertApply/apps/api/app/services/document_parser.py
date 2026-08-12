from io import BytesIO
from pathlib import Path
import re
import zipfile

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MIN_EXTRACTED_TEXT_CHARS = 100
NO_TEXT_ERROR = (
    "We could not extract text from this file. Please upload a text-based PDF/DOCX "
    "or paste your resume text."
)

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentParserError(ValueError):
    pass


def validate_uploaded_file(file: UploadFile, content: bytes | None = None) -> None:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    content_type = (file.content_type or "").split(";", 1)[0].strip().lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentParserError("Unsupported file extension. Upload a PDF, DOCX, or TXT file.")
    if content_type and content_type not in ALLOWED_CONTENT_TYPES:
        raise DocumentParserError("Unsupported file type. Upload a PDF, DOCX, or TXT file.")
    if content is not None:
        if not content:
            raise DocumentParserError("Uploaded file is empty.")
        if len(content) > MAX_UPLOAD_BYTES:
            raise DocumentParserError("Uploaded file is too large. The limit is 5MB.")


def extract_text_from_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentParserError(NO_TEXT_ERROR) from exc
    return normalize_extracted_text(text)


def extract_text_from_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text
        ]
    except (ValueError, zipfile.BadZipFile, KeyError) as exc:
        raise DocumentParserError(NO_TEXT_ERROR) from exc
    return normalize_extracted_text("\n".join([*paragraphs, *table_cells]))


# Characters that PDF extraction commonly mangles, mapped to clean equivalents.
_CHAR_REPLACEMENTS = {
    "‣": "•",  # normalize assorted bullet glyphs to one form
    "▪": "•",
    "◦": "•",
    "⁃": "•",
    "": "•",  # Symbol-font bullet from Word/PDF
    "‘": "'",  # smart quotes
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",  # en dash
    "—": "-",  # em dash
    "−": "-",  # minus sign
    " ": " ",  # non-breaking space
    "​": "",  # zero-width space
    "‌": "",  # zero-width non-joiner
    "‍": "",  # zero-width joiner
    "﻿": "",  # BOM
    "�": "",  # replacement character (hidden artifacts)
    "ﬁ": "fi",  # ligatures
    "ﬂ": "fl",
}

# Standalone page-number / artifact lines like "Page 1 of 2" or a bare "2".
_PAGE_ARTIFACT_RE = re.compile(r"^(page\s+\d+(\s+of\s+\d+)?|\d{1,3})$", re.IGNORECASE)


def normalize_extracted_text(text: str) -> str:
    for source, target in _CHAR_REPLACEMENTS.items():
        text = text.replace(source, target)
    # Re-join words broken across line boundaries by hyphenation ("infer-\nence").
    text = re.sub(r"([A-Za-z])-\n([a-z])", r"\1\2", text)

    normalized_lines: list[str] = []
    for line in text.splitlines():
        collapsed = re.sub(r"[ \t]+", " ", line).strip()
        if not collapsed:
            normalized_lines.append("")
            continue
        if _PAGE_ARTIFACT_RE.match(collapsed):
            continue  # drop page-number artifacts
        normalized_lines.append(collapsed)

    normalized = "\n".join(normalized_lines)
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def detect_document_kind(text: str, filename: str | None = None) -> str:
    lower_text = text.lower()
    lower_name = (filename or "").lower()
    if "linkedin.com/in/" in lower_text or "contact info" in lower_text and "linkedin" in lower_text:
        return "linkedin_pdf"
    if "linkedin" in lower_name:
        return "linkedin_pdf"
    if any(marker in lower_text for marker in ["experience", "education", "skills"]):
        return "resume"
    return "unknown"


def extract_uploaded_file_text(file: UploadFile, content: bytes) -> str:
    validate_uploaded_file(file, content)
    extension = Path(file.filename or "").suffix.lower()
    if extension == ".pdf":
        text = extract_text_from_pdf(content)
    elif extension == ".docx":
        text = extract_text_from_docx(content)
    elif extension == ".txt":
        try:
            text = normalize_extracted_text(content.decode("utf-8"))
        except UnicodeDecodeError as exc:
            raise DocumentParserError("Text files must use UTF-8 encoding.") from exc
    else:
        raise DocumentParserError("Unsupported file extension. Upload a PDF, DOCX, or TXT file.")
    if len(text) < MIN_EXTRACTED_TEXT_CHARS:
        raise DocumentParserError(NO_TEXT_ERROR)
    return text
